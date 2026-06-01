from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 색상 상수 ─────────────────────────────────────────────
BLUE  = RGBColor(0x00, 0x4A, 0x99)   # 충북대 상징색
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── 헬퍼 함수 ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    return p

def body(text, bold_part=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 62)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p

# ════════════════════════════════════════════════════════════
# 표지 (Page 1)
# ════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = cover_title.add_run('충 북 대 학 교')
rt.bold = True
rt.font.size = Pt(36)
rt.font.color.rgb = BLUE

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = cover_sub.add_run('Chungbuk National University')
rs.font.size = Pt(16)
rs.font.color.rgb = GRAY

doc.add_paragraph()

cover_tag = doc.add_paragraph()
cover_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
rtag = cover_tag.add_run('대학 소개 보고서')
rtag.bold = True
rtag.font.size = Pt(22)
rtag.font.color.rgb = BLACK

doc.add_paragraph()

cover_desc = doc.add_paragraph()
cover_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
rd = cover_desc.add_run('충청북도 청주시 서원구 충대로 1\n설립연도: 1951년  |  유형: 국립대학교  |  학생 수: 약 20,000명')
rd.font.size = Pt(11)
rd.font.color.rgb = GRAY

for _ in range(8):
    doc.add_paragraph()

cover_date = doc.add_paragraph()
cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
rdate = cover_date.add_run('2026년 6월')
rdate.font.size = Pt(11)
rdate.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 목차 (Page 2)
# ════════════════════════════════════════════════════════════
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = toc_title.add_run('목  차')
rt2.bold = True
rt2.font.size = Pt(20)
rt2.font.color.rgb = BLUE

doc.add_paragraph()

toc_items = [
    ('1', '대학 개요'),
    ('2', '역사 및 연혁'),
    ('3', '캠퍼스 및 시설'),
    ('4', '학사 구성 – 단과대학 및 학과'),
    ('5', '연구 역량 및 성과'),
    ('6', '국제화 및 글로벌 협력'),
    ('7', '산학협력 및 취업 지원'),
    ('8', '학생 생활 및 복지'),
    ('9', '사회 공헌 및 지역사회 연계'),
    ('10', '비전 및 발전 계획'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(2)
    run_num = p.add_run(f'{num}장.  ')
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = BLUE
    run_title = p.add_run(title)
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1장. 대학 개요 (Page 3)
# ════════════════════════════════════════════════════════════
heading1('제 1장. 대학 개요')
divider()

body(
    '충북대학교(忠北大學校, Chungbuk National University, 약칭 CBNU)는 충청북도 청주시 서원구에 위치한 '
    '대한민국의 국립종합대학교이다. 1951년 개교 이래 70여 년의 역사를 거치며 지역 거점 국립대학으로서 '
    '학술 연구와 인재 양성을 선도해 왔다.'
)

heading2('1.1 기본 현황')

# 개요 표
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data = [
    ('구  분', '내  용'),
    ('대학명',   '충북대학교 (Chungbuk National University)'),
    ('설립유형', '국립대학교'),
    ('설립연도', '1951년'),
    ('소재지',   '충청북도 청주시 서원구 충대로 1'),
    ('총장',     '정성택 (제21대)'),
    ('재학생 수', '약 20,000명 (학부 + 대학원)'),
    ('교직원 수', '약 1,200명'),
]

for i, (label, value) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    if i == 0:
        set_cell_bg(row.cells[0], '004A99')
        set_cell_bg(row.cells[1], '004A99')
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = WHITE
    else:
        set_cell_bg(row.cells[0], 'E8F0FB')
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = BLUE

doc.add_paragraph()
body(
    '충북대학교는 "진리·정의·개척"을 교훈으로, 지역과 국가를 넘어 세계적 수준의 연구중심대학으로 도약하기 위해 '
    '지속적으로 노력하고 있다. 현재 13개 단과대학, 70여 개 학과(부), 대학원 등으로 구성된 종합대학으로서 '
    '다양한 학문 분야를 아우르고 있다.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2장. 역사 및 연혁 (Page 4)
# ════════════════════════════════════════════════════════════
heading1('제 2장. 역사 및 연혁')
divider()

body(
    '충북대학교는 광복 이후 충청북도의 고등교육 수요에 부응하기 위해 설립되었다. '
    '이후 국가와 지역사회의 발전과 함께 성장하며 명실상부한 거점 국립대학으로 자리매김하였다.'
)

heading2('2.1 주요 연혁')

timeline = [
    ('1951', '충북대학교 개교 (농과대학·공과대학 중심)'),
    ('1957', '종합대학교 승격'),
    ('1970년대', '이공계 특성화 및 단과대학 확장'),
    ('1980년대', '법과대학·사범대학 신설, 대학원 확대'),
    ('1990년대', '정보·전자 분야 학과 신설, 국제화 추진 시작'),
    ('2000년대', 'BK21 사업 선정, 연구중심대학 도약'),
    ('2010년대', '산학협력 선도대학(LINC) 선정, 글로벌 협력 확대'),
    ('2020년대', '디지털 전환·AI 특성화, 혁신공유대학 참여'),
]

for year, event in timeline:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    yr = p.add_run(f'▶ {year}  ')
    yr.bold = True
    yr.font.size = Pt(10.5)
    yr.font.color.rgb = BLUE
    ev = p.add_run(event)
    ev.font.size = Pt(10.5)
    ev.font.color.rgb = GRAY

body(
    '현재 충북대학교는 개교 70주년을 맞아 "세계를 향한 충북대, 충북을 이끄는 충북대"라는 슬로건 아래 '
    '혁신적인 교육 및 연구 환경 구축에 매진하고 있다.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3장. 캠퍼스 및 시설 (Page 5)
# ════════════════════════════════════════════════════════════
heading1('제 3장. 캠퍼스 및 시설')
divider()

body(
    '충북대학교 메인 캠퍼스는 약 206만 m²(63만 평)에 달하는 넓은 부지 위에 조성되어 있으며, '
    '오송·증평 등 분교 캠퍼스도 운영 중이다. 자연 친화적인 환경과 현대적 교육시설이 조화를 이루고 있다.'
)

heading2('3.1 주요 시설')

facilities = [
    ('학술정보원(도서관)', '장서 140만 권 이상, 전자 자료·열람실·스터디룸 완비'),
    ('첨단강의동·스마트강의실', '최신 ICT 기자재 탑재, 원격·혼합 수업 가능'),
    ('학생회관 및 복지시설', '식당·카페·은행·우체국·의원 등 생활 인프라'),
    ('기숙사(생활관)',   '학부·대학원생 약 3,000명 수용 규모'),
    ('스포츠센터·운동장', '체육관·수영장·테니스장·다목적 구장'),
    ('창업지원센터·산학협력관', '스타트업 인큐베이팅, 기업 협업 공간'),
    ('오송 바이오캠퍼스', '의약·바이오 분야 연구시설 및 임상 공간'),
]

for name, desc in facilities:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    rn = p.add_run(f'● {name}  ')
    rn.bold = True
    rn.font.size = Pt(10.5)
    rn.font.color.rgb = BLUE
    rd2 = p.add_run(f'— {desc}')
    rd2.font.size = Pt(10.5)
    rd2.font.color.rgb = GRAY

doc.add_paragraph()
body(
    '최근에는 제로 에너지 건물 건축, 태양광 패널 설치 등 탄소중립 캠퍼스 조성 사업도 적극 추진 중이며, '
    '장애학생 편의시설 확충과 배리어프리(Barrier-free) 캠퍼스 구현에도 힘쓰고 있다.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4장. 학사 구성 (Page 6)
# ════════════════════════════════════════════════════════════
heading1('제 4장. 학사 구성 – 단과대학 및 학과')
divider()

body(
    '충북대학교는 인문·사회·이공·예체능 전 분야를 포괄하는 13개 단과대학 체제로 운영되며, '
    '대학원(일반대학원·전문대학원·특수대학원)을 통해 석·박사 과정도 운영한다.'
)

heading2('4.1 단과대학 현황')

colleges = [
    ('인문대학',     '국어국문학과, 영어영문학과, 사학과, 철학과 등'),
    ('사회과학대학', '사회학과, 심리학과, 행정학과, 정치외교학과 등'),
    ('경영대학',     '경영학부, 경제학과, 회계학과 등'),
    ('법과대학',     '법학부'),
    ('자연과학대학', '수학과, 물리학과, 화학과, 생물학과 등'),
    ('공과대학',     '기계공학부, 전기공학부, 화학공학과, 건축공학과 등'),
    ('전자정보대학', '전자공학부, 정보통신공학부, 소프트웨어학부 등'),
    ('농업생명환경대학', '식물자원학과, 축산학과, 환경공학과 등'),
    ('수의과대학',   '수의학과 (6년제)'),
    ('의과대학',     '의학과 (6년제)'),
    ('약학대학',     '약학과 (6년제)'),
    ('사범대학',     '교육학과, 수학교육과, 과학교육과, 체육교육과 등'),
    ('예술대학',     '음악학과, 미술학과, 디자인학과 등'),
]

table2 = doc.add_table(rows=len(colleges)+1, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table2.rows[0]
hdr.cells[0].text = '단과대학'
hdr.cells[1].text = '주요 학과'
set_cell_bg(hdr.cells[0], '004A99')
set_cell_bg(hdr.cells[1], '004A99')
for cell in hdr.cells:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE

for i, (college, depts) in enumerate(colleges):
    row = table2.rows[i+1]
    row.cells[0].text = college
    row.cells[1].text = depts
    set_cell_bg(row.cells[0], 'E8F0FB')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = BLUE

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5장. 연구 역량 (Page 7)
# ════════════════════════════════════════════════════════════
heading1('제 5장. 연구 역량 및 성과')
divider()

body(
    '충북대학교는 국가 R&D 과제 수주 규모를 지속적으로 확대하며 연구중심대학으로의 전환을 가속화하고 있다. '
    'BK21·4단계 사업단 다수 선정, SCI(E)급 논문 게재 실적 증가, 특허 출원·등록 확대 등이 주요 성과이다.'
)

heading2('5.1 주요 연구 성과 지표')

metrics = [
    ('연간 연구비 수주액', '약 2,500억 원 (2024년 기준)'),
    ('SCI(E) 논문 게재 수', '연간 2,000편 이상'),
    ('특허 출원·등록',     '연간 500건 이상'),
    ('BK21 선정 사업단',   '10개 사업단 (이공계·인문사회 포함)'),
    ('국책 연구센터',       '20개 이상 (정부·민간 지정)'),
]

for key, val in metrics:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    rk = p.add_run(f'◆ {key}: ')
    rk.bold = True
    rk.font.size = Pt(10.5)
    rk.font.color.rgb = BLUE
    rv = p.add_run(val)
    rv.font.size = Pt(10.5)
    rv.font.color.rgb = GRAY

heading2('5.2 특화 연구 분야')
for area in ['반도체·디스플레이 소재', '바이오·의약 분야 (오송 바이오클러스터 연계)',
             '스마트 농업·환경 기술', 'AI·빅데이터·사이버보안', '탄소중립·신재생에너지']:
    bullet(area)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6장. 국제화 (Page 8)
# ════════════════════════════════════════════════════════════
heading1('제 6장. 국제화 및 글로벌 협력')
divider()

body(
    '충북대학교는 50개국 이상 200여 개 해외 대학·기관과 교류협정을 체결하고, '
    '교환학생·공동학위·글로벌 인턴십 등 다양한 국제화 프로그램을 운영하고 있다.'
)

heading2('6.1 주요 국제 협력 현황')
for item in [
    '교류협정 대학: 50개국 200여 개교',
    '외국인 유학생: 약 1,500명 (학위·어학연수 포함)',
    '교환학생 파견·수용: 연간 500명 이상',
    '영어 강의(EMI) 비율: 전체 강좌의 20% 이상',
    '글로벌 현장실습(해외 인턴십): 연간 100명 이상 파견',
]:
    bullet(item)

heading2('6.2 주요 협력 국가 및 대학')
for item in [
    '미국: University of Illinois, Purdue University 등',
    '중국: 베이징대학교, 상하이교통대학교 등',
    '일본: 오사카대학교, 규슈대학교 등',
    '유럽: 독일·프랑스·스페인 주요 공과대학',
    '동남아: 말레이시아·베트남·인도네시아 거점대학',
]:
    bullet(item)

body(
    '최근에는 QS 세계대학평가 참여 및 글로벌 랭킹 상승을 목표로 연구 역량 강화와 국제 공동연구를 확대하고 있다.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7장. 산학협력 & 8장. 학생생활 (Page 9)
# ════════════════════════════════════════════════════════════
heading1('제 7장. 산학협력 및 취업 지원')
divider()

body(
    '충북대학교는 LINC(산학협력 선도대학) 사업 연속 선정, 테크노파크 및 오창·오송 산업단지와의 긴밀한 협력을 통해 '
    '실무형 인재 양성과 기술 이전·창업 활성화를 추진하고 있다.'
)

heading2('7.1 주요 산학협력 프로그램')
for item in [
    '캡스톤 디자인: 기업 연계 프로젝트 수업 전 학과 운영',
    '현장실습(인턴십): 연간 1,000명 이상 국내·외 기업 파견',
    '창업 교육 및 지원: 창업 동아리 100개 이상, 창업 공간 제공',
    '기술이전 및 공동연구: 지역 중소·중견기업과 연간 200건 이상 계약',
    '취업률: 졸업생 취업률 65~70% 수준 유지 (국립대 상위권)',
]:
    bullet(item)

doc.add_paragraph()

heading1('제 8장. 학생 생활 및 복지')
divider()

body(
    '충북대학교는 재학생의 학업·생활 전반을 지원하기 위해 다양한 장학금 제도와 복지 프로그램을 운영한다.'
)

heading2('8.1 장학금 제도')
for item in [
    '국가장학금(한국장학재단) 연계: 소득분위별 지원',
    '교내 성적 우수 장학금: 상위 10% 자동 선발',
    '근로 장학금·긴급 생활비 지원',
    '외국어·자격증 취득 장학금',
]:
    bullet(item)

heading2('8.2 학생 지원 서비스')
for item in [
    '심리상담센터: 개인·집단 상담, 심리검사',
    '학생건강센터: 기초 의료 서비스 무료 제공',
    '커리어개발센터: 취업·진로 상담, 모의면접',
    '다문화·외국인 학생 전담 지원팀',
]:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9장. 사회공헌 & 10장. 비전 (Page 10)
# ════════════════════════════════════════════════════════════
heading1('제 9장. 사회 공헌 및 지역사회 연계')
divider()

body(
    '충북대학교는 충청북도 및 청주시와의 협력을 통해 지역 현안 해결, 평생교육 제공, '
    '문화·예술 행사 개최 등 다양한 사회 공헌 활동을 펼치고 있다.'
)

for item in [
    '지역 맞춤형 인재 양성 협약(충청북도·청주시)',
    '지역 중소기업 기술 컨설팅 무료 제공',
    '평생교육원 운영: 일반 시민 대상 200여 강좌',
    '농어촌 봉사활동·사회적 약자 지원 프로그램',
    '충북 지역 문화재 발굴·보존 연구 참여',
]:
    bullet(item)

doc.add_paragraph()

heading1('제 10장. 비전 및 발전 계획')
divider()

body(
    '충북대학교는 "글로벌 TOP100 거점 국립대학"을 중장기 목표로 설정하고, '
    '교육·연구·산학·국제화 전 분야에서 혁신적 변화를 추진하고 있다.'
)

heading2('10.1 4대 전략 방향')
strategies = [
    ('교육 혁신',   '학생 중심 교육과정 개편, AI·디지털 역량 강화, PBL 확대'),
    ('연구 고도화', '대형 연구과제 수주, 글로벌 공동연구, 스타 연구자 유치'),
    ('산학 협력',   '지역 산업 맞춤형 인재 공급, 딥테크 스타트업 육성'),
    ('글로벌화',    '영어 강의 확대, 외국인 교원·학생 유치, 해외 캠퍼스 거점 설립'),
]

for i, (name, desc) in enumerate(strategies):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(5)
    rn2 = p.add_run(f'전략 {i+1}. {name}  ')
    rn2.bold = True
    rn2.font.size = Pt(11)
    rn2.font.color.rgb = BLUE
    rd3 = p.add_run(f'— {desc}')
    rd3.font.size = Pt(10.5)
    rd3.font.color.rgb = GRAY

doc.add_paragraph()
body(
    '충북대학교는 앞으로도 "진리·정의·개척"의 교훈 아래 지역을 넘어 세계와 함께하는 대학으로 성장하며, '
    '대한민국 고등교육 발전과 지역사회 번영에 기여할 것이다.'
)

doc.add_paragraph()
divider()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_r = end_p.add_run('— 이 상 —')
end_r.bold = True
end_r.font.size = Pt(12)
end_r.font.color.rgb = BLUE

# ── 저장 ──────────────────────────────────────────────────
output_path = r'C:\Users\LG\Desktop\충북대학교_소개_보고서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
